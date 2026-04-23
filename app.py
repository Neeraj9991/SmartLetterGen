import os
import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
import zipfile
import re

st.set_page_config(page_title="SGV Group | Smart Letter Gen", page_icon="🔴", layout="centered")

# ===================== CUSTOM CSS =====================
st.markdown("""
    <style>
        .stApp {
            background-color: #F9FAFB;
        }
        /* Custom Header Banner */
        .custom-banner {
            background-color: #264653;
            padding: 24px;
            border-radius: 12px;
            display: flex;
            align-items: center;
            gap: 20px;
            margin-bottom: 30px;
            box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
        }
        .logo-box {
            background-color: white;
            padding: 8px;
            border-radius: 10px;
            width: 80px;
            height: 80px;
            display: flex;
            justify-content: center;
            align-items: center;
            flex-shrink: 0;
        }    
        .banner-title {
            margin: 0 !important;
            font-size: 1.8rem !important;
            font-weight: 700 !important;
            color: white !important;
        }
        .banner-subtitle {
            margin: 5px 0 0 0;
            font-size: 0.95rem;
            color: #E2E8F0;
        }
        /* Section Titles */
        .section-title {
            color: #64748B;
            font-size: 0.9rem;
            font-weight: 700;
            letter-spacing: 1px;
            margin-top: 10px;
            margin-bottom: 10px;
            text-transform: uppercase;
        }
        /* Uploader Boxes */
        div[data-testid="stFileUploader"] {
            padding: 1.5rem;
            border-radius: 10px;
            background-color: #F1F5F9;
            border: 2px dashed #CBD5E1;
            transition: all 0.3s;
        }
        div[data-testid="stFileUploader"]:hover {
            border-color: #94A3B8;
            background-color: #E2E8F0;
        }
        /* Button */
        .stButton>button {
            width: 100%;
            border-radius: 8px;
            font-weight: bold;
            transition: all 0.3s;
        }
        .stButton>button:hover {
            transform: translateY(-2px);
            box-shadow: 0 4px 6px rgba(49, 130, 206, 0.2);
        }
    </style>
""", unsafe_allow_html=True)

# ===================== SIDEBAR =====================
st.sidebar.title("Excel Rules")

st.sidebar.markdown("""
### Mandatory Excel Columns
- **1st Column:** ECode
- **2nd Column:** Name

---

### Styling Suffixes
- `_b` → Bold  
- `_i` → Italic  
- `_u` → Underline  

You can combine them:
- `_bi`, `_bu`, `_iu`, `_biu`

---

### Number Formatting
- `_c` → Currency (₹)  
- `_comma` → Number with commas  
- `_2d` → 2 decimal places  

---

### Examples
- `Salary_c_b` → **₹5,000**  
- `Score_2d_i` → *98.45*  
- `Users_comma_b` → **1,000,000**  

No suffix = plain text
""")

# ---------- SAMPLE EXCEL ----------
sample_df = pd.DataFrame({
    "ECode": ["EMP001"],
    "Name_b": ["Neeraj Balodi"],
    "Designation_i": ["Manager"],
    "Department": ["IT"],
    "Salary_c_b": [5000],
    "Score_2d_i": [98.456],
    "Users_comma_b": [1000000]
})

sample_buffer = BytesIO()
sample_df.to_excel(sample_buffer, index=False, engine="openpyxl")
sample_buffer.seek(0)

st.sidebar.download_button(
    label="⬇️ Download Sample Excel",
    data=sample_buffer,
    file_name="SmartLetterGen_Sample.xlsx",
    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    type="primary"
)

# ===================== MAIN =====================
import os
import base64

# Helper to load logo as base64
logo_b64 = ""
if os.path.exists("logo.png"):
    with open("logo.png", "rb") as f:
        logo_b64 = base64.b64encode(f.read()).decode()

logo_html = f'<img src="data:image/png;base64,{logo_b64}" style="max-width: 100%; max-height: 100%;">' if logo_b64 else '<div style="color: #B91C1C; font-size: 22px; font-weight: 900; text-align: center; letter-spacing: 1px;">SGV</div>'

st.markdown(f"""
    <div class="custom-banner">
        <div class="logo-box">
            {logo_html}
        </div>
        <div>
            <h1 class="banner-title">Smart Letter Gen</h1>
            <p class="banner-subtitle">Upload your DOCX template and Excel data — personalized letters will be generated automatically.</p>
        </div>
    </div>
""", unsafe_allow_html=True)

st.markdown('<p class="section-title">📁 UPLOAD FILES</p>', unsafe_allow_html=True)

col_t, col_e = st.columns(2)
with col_t:
    template_file = st.file_uploader("DOCX Template", type=["docx"])
with col_e:
    excel_file = st.file_uploader("Excel Data (.xlsx)", type=["xlsx"])

# ===================== FORMATTERS =====================
def format_value(key, value):
    if value is None:
        return ""

    # Remove decimal .0 for whole numbers parsed as floats by pandas
    if type(value).__name__ in ['float', 'float32', 'float64'] and value == int(value):
        value_str = str(int(value))
    else:
        value_str = str(value)

    key_lower = key.lower()

    # Comma formatting (check before _c since "_comma" contains "_c")
    if "_comma" in key_lower:
        try:
            value_str = f"{int(float(value)):,}"
        except:
            pass

    # Currency
    elif "_c" in key_lower:
        try:
            value_str = f"₹{int(float(value)):,}"
        except:
            pass

    # 2 decimal formatting
    elif "_2d" in key_lower:
        try:
            value_str = f"{float(value):.2f}"
        except:
            pass

    return value_str


def get_text_style(key):
    style = {"bold": False, "italic": False, "underline": False}
    key_lower = key.lower()

    if "_b" in key_lower:
        style["bold"] = True
    if "_i" in key_lower:
        style["italic"] = True
    if "_u" in key_lower:
        style["underline"] = True

    return style

# ===================== PLACEHOLDER ENGINE =====================
def replace_placeholders(doc, data):
    pattern = re.compile(r"\{\s*(.*?)\s*\}")

    def process_paragraph(para):
        full_text = "".join(run.text for run in para.runs)
        if not full_text:
            return

        matches = pattern.findall(full_text)
        if not matches:
            return

        # Preserve original font size from the first run
        original_font_size = None
        if para.runs:
            original_font_size = para.runs[0].font.size

        para.clear()
        last_idx = 0

        for match in matches:
            placeholder = f"{{{match}}}"
            start = full_text.find(placeholder, last_idx)

            if start > last_idx:
                run = para.add_run(full_text[last_idx:start])
                if original_font_size:
                    run.font.size = original_font_size

            key = match.strip()
            if key in data:
                raw_value = data[key]
                formatted_value = format_value(key, raw_value)
                style = get_text_style(key)

                run = para.add_run(formatted_value)
                run.bold = style["bold"]
                run.italic = style["italic"]
                run.underline = style["underline"]
                # Preserve original font size
                if original_font_size:
                    run.font.size = original_font_size

            last_idx = start + len(placeholder)

        if last_idx < len(full_text):
            run = para.add_run(full_text[last_idx:])
            if original_font_size:
                run.font.size = original_font_size

    # Paragraphs
    for para in doc.paragraphs:
        process_paragraph(para)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_paragraph(para)

# ===================== PROCESS =====================
if template_file and excel_file:
    df = pd.read_excel(excel_file)

    # ---------- VALIDATION ----------
    if df.shape[1] < 2:
        st.error("❌ Excel must contain at least ECode and Name columns.")
        st.stop()

    if not df.columns[0].lower().startswith("ecode"):
        st.error("❌ First column must start with 'ECode'.")
        st.stop()

    if not df.columns[1].lower().startswith("name"):
        st.error("❌ Second column must start with 'Name'.")
        st.stop()

    st.success("✅ Excel structure validated successfully")

    st.subheader("📊 Detected Excel Headers")
    st.write(list(df.columns))

    if st.button("🚀 Generate Letters (ZIP)", type="primary"):
        zip_buffer = BytesIO()

        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            for _, row in df.iterrows():
                doc = Document(template_file)

                row_data = {
                    col: "" if pd.isna(row[col]) else row[col]
                    for col in df.columns
                }

                replace_placeholders(doc, row_data)

                doc_buffer = BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)

                filename = f"{row[df.columns[0]]}_{row[df.columns[1]]}.docx"
                zip_file.writestr(filename, doc_buffer.read())

        zip_buffer.seek(0)

        st.download_button(
            label="⬇️ Download Letters (ZIP)",
            data=zip_buffer,
            file_name="SmartLetterGen_Output.zip",
            mime="application/zip",
            type="primary"
        )
else:
    st.markdown("<br>", unsafe_allow_html=True)
    st.info("👆 Upload both your DOCX template and Excel data to get started.")
